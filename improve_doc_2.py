
<<<<<<< HEAD
This script converts a large DOCX document to a new template format while
preserving structural elements using Claude via Amazon Bedrock.
"""

import docx
from tqdm import tqdm
import argparse
import boto3
import json
import os
from concurrent.futures import ThreadPoolExecutor

# Set up argument parser
parser = argparse.ArgumentParser(description='Convert large DOCX document to a new template structure using Claude')
parser.add_argument('--original', required=True, help='Path to the original document')
parser.add_argument('--template', required=True, help='Path to the new template document')
parser.add_argument('--output', required=True, help='Path to save the reformatted document')
parser.add_argument('--aws-profile', help='AWS profile to use (optional if using default profile)')
parser.add_argument('--aws-region', default='us-east-1', help='AWS region where Bedrock is available')
parser.add_argument('--model', default='anthropic.claude-3-sonnet-20240229-v1:0', help='Bedrock model ID to use')
parser.add_argument('--chunk-size', type=int, default=10, help='Number of paragraphs per chunk')
parser.add_argument('--max-workers', type=int, default=4, help='Number of parallel workers for processing chunks')

def extract_document_structure(doc_path):
    """Extract structure from a docx file including headers, subheaders, and paragraphs."""
    print(f"Extracting structure from {doc_path}")
    doc = docx.Document(doc_path)
    structure = []
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue  # Skip empty paragraphs
            
        # Determine paragraph type based on style
        para_type = "paragraph"
        if para.style.name.startswith("Heading 1") or para.style.name.lower() == "heading 1":
            para_type = "header"
        elif para.style.name.startswith("Heading 2") or para.style.name.lower() == "heading 2":
            para_type = "subheader"
        
        # Extract formatting properties (simplified)
        formatting = {
            "style_name": para.style.name
        }
        
        structure.append({
            "type": para_type,
            "text": para.text,
            "formatting": formatting
        })
    
    print(f"Extracted {len(structure)} paragraphs from {doc_path}")
    return structure

import json

## template reading json file 
def analyze_template_structure(json_template_path):
    """Analyze JSON template to extract formatting rules for different elements."""
    print("Analyzing JSON template structure")
    formatting_rules = {
        "header": None,
        "subheader": None,
        "paragraph": None
    }

    # Load JSON template
    with open(json_template_path, 'r') as f:
        template_data = json.load(f)

    # Extract formatting rules from JSON structure
    for item in template_data:
        element_type = item.get("type")
        if element_type in formatting_rules and formatting_rules[element_type] is None:
            formatting_rules[element_type] = item.get("formatting", {}).get("style_name")

    # Set defaults for any missing elements
    formatting_rules.setdefault("header", "Heading 1")
    formatting_rules.setdefault("subheader", "Heading 2") 
    formatting_rules.setdefault("paragraph", "Normal")

    return formatting_rules

## template reading docx file 
def analyze_template_structure(template_structure):
    """Analyze the template to extract formatting rules for different elements."""
    print("Analyzing template structure")
    formatting_rules = {
        "header": None,
        "subheader": None,
        "paragraph": None
    }
    
    for item in template_structure:
        if item["type"] in formatting_rules and formatting_rules[item["type"]] is None:
            formatting_rules[item["type"]] = item["formatting"]["style_name"]
    
    # If we didn't find examples of all types, use defaults
    if formatting_rules["header"] is None:
        formatting_rules["header"] = "Heading 1"
    if formatting_rules["subheader"] is None:
        formatting_rules["subheader"] = "Heading 2" 
    if formatting_rules["paragraph"] is None:
        formatting_rules["paragraph"] = "Normal"
    
    return formatting_rules

def chunk_document(structure, chunk_size=10):
    """Break down document into chunks based on headers."""
    print(f"Breaking document into chunks (chunk size: {chunk_size})")
    chunks = []
    current_chunk = []
    current_header = None
    current_subheader = None
    
    for item in structure:
        if item["type"] == "header":
            # If we already have content and find a new header, save the chunk
            if current_chunk:
                chunks.append({
                    "header": current_header,
                    "subheader": current_subheader,
                    "content": current_chunk.copy()
                })
            
            # Start a new chunk
            current_header = item
            current_subheader = None
            current_chunk = []
        
        elif item["type"] == "subheader":
            # If we find a new subheader and the current chunk is large, save it
            if current_chunk and len(current_chunk) >= chunk_size:
                chunks.append({
                    "header": current_header,
                    "subheader": current_subheader,
                    "content": current_chunk.copy()
                })
                current_chunk = []
            
            current_subheader = item
            current_chunk.append(item)
        
        else:  # Regular paragraph
            current_chunk.append(item)
            
            # If chunk reaches size limit and we're not in the middle of a section,
            # save it and start a new one
            if len(current_chunk) >= chunk_size:
                chunks.append({
                    "header": current_header,
                    "subheader": current_subheader,
                    "content": current_chunk.copy()
                })
                current_chunk = []
    
    # Add the last chunk if there's anything left
    if current_chunk:
        chunks.append({
            "header": current_header,
            "subheader": current_subheader,
            "content": current_chunk.copy()
        })
    
    print(f"Document broken into {len(chunks)} chunks")
    return chunks

def apply_template_with_claude(chunk, template_rules, bedrock_client, model_id):
    """Use Anthropic Claude on Amazon Bedrock to reformat the chunk."""
    # Prepare the content for the LLM
    chunk_text = ""
    
    if chunk["header"]:
        chunk_text += f"# Header: {chunk['header']['text']}\n\n"
    
    if chunk["subheader"]:
        chunk_text += f"## Subheader: {chunk['subheader']['text']}\n\n"
    
    for item in chunk["content"]:
        if item["type"] == "paragraph":
            chunk_text += f"{item['text']}\n\n"
    
    # Create a prompt for Claude
    prompt = f"""
    I have a document section that needs to be reformatted according to a new template.
    
    The new template has these formatting rules:
    - Headers: {template_rules['header']}
    - Subheaders: {template_rules['subheader']}
    - Paragraphs: {template_rules['paragraph']}
    
    Here's the content to reformat:
    
    {chunk_text}
    
    Please reformat this content according to the template rules, preserving all original information.
    Return the reformatted content with clear separation between headers, subheaders, and paragraphs.
    
    Format your response like this:
    # HEADER: <header text>
    ## SUBHEADER: <subheader text>
    PARAGRAPH: <paragraph text>
    
    (Only include header/subheader if present in the original)
    """
    
    # Create the request body for Claude
    request_body = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 1500,
        "temperature": 0.2,
        "system": "You are a document formatting assistant.",
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    }
    
    try:
        # Call the Bedrock API
        response = bedrock_client.invoke_model(
            modelId=model_id,
            body=json.dumps(request_body)
        )
        
        # Parse the response
        response_body = json.loads(response["body"].read())
        reformatted_text = response_body["content"][0]["text"]
        
        # Parse the reformatted content back into structured form
        reformatted_structure = parse_llm_response(reformatted_text, template_rules)
        
        return reformatted_structure
    
    except Exception as e:
        print(f"Error processing chunk with Claude: {e}")
        # Return the original chunk structure as fallback
        if chunk["subheader"]:
            return [chunk["header"]] + [chunk["subheader"]] + chunk["content"]
        else:
            return [chunk["header"]] + chunk["content"]

def parse_llm_response(response_text, template_rules):
    """Parse the LLM's response back into structured form."""
    lines = response_text.split("\n")
    structure = []
    current_type = None
    current_text = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith("# HEADER:"):
            # Save any previously gathered text
            if current_type and current_text:
                structure.append({
                    "type": current_type,
                    "text": " ".join(current_text),
                    "formatting": {"style_name": template_rules[current_type]}
                })
                current_text = []
            
            # Start a new header
            current_type = "header"
            current_text = [line.replace("# HEADER:", "").strip()]
            
        elif line.startswith("## SUBHEADER:"):
            # Save any previously gathered text
            if current_type and current_text:
                structure.append({
                    "type": current_type,
                    "text": " ".join(current_text),
                    "formatting": {"style_name": template_rules[current_type]}
                })
                current_text = []
            
            # Start a new subheader
            current_type = "subheader"
            current_text = [line.replace("## SUBHEADER:", "").strip()]
            
        elif line.startswith("PARAGRAPH:"):
            # Save any previously gathered text
            if current_type and current_text:
                structure.append({
                    "type": current_type,
                    "text": " ".join(current_text),
                    "formatting": {"style_name": template_rules[current_type]}
                })
                current_text = []
            
            # Start a new paragraph
            current_type = "paragraph"
            current_text = [line.replace("PARAGRAPH:", "").strip()]
            
        else:
            # Continue current text
            if current_type:
                current_text.append(line)
    
    # Save the last paragraph
    if current_type and current_text:
        structure.append({
            "type": current_type,
            "text": " ".join(current_text),
            "formatting": {"style_name": template_rules[current_type]}
        })
    
    return structure

def process_chunks_parallel(chunks, template_rules, bedrock_client, model_id, max_workers):
    """Process chunks in parallel using ThreadPoolExecutor."""
    print(f"Processing {len(chunks)} chunks with {max_workers} workers")
    reformatted_chunks = []
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = []
        for chunk in chunks:
            future = executor.submit(
                apply_template_with_claude, 
                chunk, 
                template_rules, 
                bedrock_client,
                model_id
            )
            futures.append(future)
        
        # Process results as they complete
        for i, future in enumerate(tqdm(futures, desc="Processing chunks with Claude")):
            try:
                result = future.result()
                reformatted_chunks.append(result)
            except Exception as e:
                print(f"Error processing chunk {i}: {e}")
    
    return reformatted_chunks

def create_new_document(reformatted_chunks, output_path):
    """Create a new document with the reformatted content."""
    print(f"Creating new document at {output_path}")
    
    # Start with a new document
    doc = docx.Document()
    
    # Flatten the chunks
    flattened_structure = []
    for chunk in reformatted_chunks:
        if isinstance(chunk, list):
            flattened_structure.extend(chunk)
        else:
            # Sometimes LLM might return a single item
            flattened_structure.append(chunk)
    
    # Add each paragraph to the document with appropriate style
    for para in flattened_structure:
        if not para or not isinstance(para, dict) or "text" not in para:
            continue  # Skip invalid items
            
        p = doc.add_paragraph(para.get("text", ""))
        
        # Apply style based on type
        if para.get("type") == "header":
            p.style = "Heading 1"
        elif para.get("type") == "subheader":
            p.style = "Heading 2"
        else:
            p.style = "Normal"
    
    # Save the document
    doc.save(output_path)
    return output_path

def main():
    """Main function to run the document conversion process."""
    args = parser.parse_args()
    
    # Set up AWS session and client
    session_args = {}
    if args.aws_profile:
        session_args['profile_name'] = args.aws_profile
    
    session = boto3.Session(**session_args, region_name=args.aws_region)
    bedrock_client = session.client('bedrock-runtime')
    
    # 1. Extract structure from both documents
    original_structure = extract_document_structure(args.original)
    template_structure = extract_document_structure(args.template)
    
    # 2. Analyze the template structure
    template_rules = analyze_template_structure(template_structure)
    print(f"Template rules: {template_rules}")
    
    # 3. Chunk the original document
    document_chunks = chunk_document(original_structure, args.chunk_size)
    
    # 4. Process chunks with Claude
    reformatted_chunks = process_chunks_parallel(
        document_chunks, 
        template_rules, 
        bedrock_client,
        args.model,
        args.max_workers
    )
    
    # 5. Create the new document
    output_path = create_new_document(reformatted_chunks, args.output)
    print(f"Document successfully reformatted and saved to {output_path}")

if __name__ == "__main__":
    main()
=======
>>>>>>> b3d7fad56d66edd03b3b6906a427b424876d88a6
